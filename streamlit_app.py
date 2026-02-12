import streamlit as st
import os
import requests
import base64
import pdfminer.high_level
import docx2txt
from PIL import Image
from pdf2image import convert_from_bytes
import io
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# --- 1. PAGE SETUP ---
st.set_page_config(page_title="Pro Exam & Biz Center", page_icon="💎", layout="wide")

if "generated_content" not in st.session_state:
    st.session_state.generated_content = ""
if "typst_code" not in st.session_state:
    st.session_state.typst_code = ""

# --- 2. SMART MODEL MANAGER ---
def get_working_model(api_key, deep_search=False):
    url = f"https://generativelanguage.googleapis.com/v1beta/models?key={api_key}"
    try:
        response = requests.get(url)
        if response.status_code == 200:
            data = response.json()
            available = [m['name'].replace('models/', '') for m in data.get('models', []) if 'generateContent' in m.get('supportedGenerationMethods', [])]
            
            if deep_search and "gemini-1.5-pro" in available: return "gemini-1.5-pro"
            if "gemini-1.5-flash" in available: return "gemini-1.5-flash"
            if "gemini-pro" in available: return "gemini-pro"
            if available: return available[0]
    except: pass
    return "gemini-pro"

# --- 3. API CALL ---
def call_gemini(api_key, model, prompt, content_parts):
    url = f"https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent?key={api_key}"
    headers = {'Content-Type': 'application/json'}
    parts = [{"text": prompt}]
    for item in content_parts:
        if item["type"] == "text": parts.append({"text": item["data"]})
        elif item["type"] == "image": parts.append({"inline_data": {"mime_type": "image/jpeg", "data": item["data"]}})
            
    data = {"contents": [{"parts": parts}], "generationConfig": {"temperature": 0.2}}
    try:
        response = requests.post(url, headers=headers, json=data)
        if response.status_code == 200: return response.json()['candidates'][0]['content']['parts'][0]['text']
        else: return f"Error: {response.text}"
    except Exception as e: return f"Connection Error: {e}"

# --- 4. FILE PROCESSOR ---
def encode_image(image):
    buffered = io.BytesIO()
    image.save(buffered, format="JPEG")
    return base64.b64encode(buffered.getvalue()).decode('utf-8')

def process_files(uploaded_files, vision_mode=False, start_page=1, end_page=None):
    content_parts = []
    if not uploaded_files: return content_parts
    for file in uploaded_files:
        try:
            ext = file.name.split('.')[-1].lower()
            if ext == 'pdf':
                if vision_mode:
                    st.toast(f"📸 Scanning {file.name}...", icon="⏳")
                    try:
                        images = convert_from_bytes(file.read(), first_page=start_page, last_page=end_page)
                        for img in images: content_parts.append({"type": "image", "data": encode_image(img)})
                    except: st.error("PDF Error: Check packages.txt")
                else:
                    text = pdfminer.high_level.extract_text(file)
                    content_parts.append({"type": "text", "data": text})
            elif ext in ['png', 'jpg', 'jpeg']:
                img = Image.open(file)
                content_parts.append({"type": "image", "data": encode_image(img)})
            elif ext == 'docx':
                text = docx2txt.process(file)
                content_parts.append({"type": "text", "data": text})
        except: pass
    return content_parts

# --- 5. WORD GENERATOR (WITH COVER PAGE) ---
def create_docx(text, is_exam_paper=False):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    
    # --- AUTOMATIC COVER PAGE (NEW) ---
    if is_exam_paper:
        heading = doc.add_heading("E-LEARNING EXAM PAPER", 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Cover Details
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run("\nSubject: Mathematics / Science\n").bold = True
        p.add_run("Grade: 10/11\n").bold = True
        p.add_run("Time: 2 Hours\n\n")
        p.add_run("______________________________________________________\n")
        p.add_run("Name: ..............................................................\n")
        p.add_run("Index No: ........................................................\n")
        p.add_run("______________________________________________________\n\n")
        doc.add_page_break()
    else:
        doc.add_heading("Generated Document", 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    lines = text.split('\n')
    table_mode = False
    table_data = []

    for line in lines:
        line = line.strip()
        if not line: continue

        if "|" in line and len(line.split("|")) > 2:
            table_mode = True
            row_data = [cell.strip() for cell in line.split("|") if cell.strip()]
            if "---" not in line: table_data.append(row_data)
            continue
        
        if table_mode and ("|" not in line):
            table_mode = False
            if table_data:
                try:
                    cols = len(table_data[0])
                    table = doc.add_table(rows=len(table_data), cols=cols)
                    table.style = 'Table Grid'
                    for i, row in enumerate(table_data):
                        for j, cell in enumerate(row):
                            if j < cols: table.cell(i, j).text = cell
                    doc.add_paragraph()
                except: pass
                table_data = []

        if is_exam_paper:
            if "MARKING SCHEME" in line: # New Page for Answers
                doc.add_page_break()
                h = doc.add_paragraph()
                h.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = h.add_run(line)
                run.bold = True
                run.font.size = Pt(14)
                run.underline = True
            elif line.startswith("#") or "Question" in line or "Part" in line:
                p = doc.add_paragraph()
                run = p.add_run(line.replace("#", "").strip())
                run.bold = True
                run.font.size = Pt(12)
                p.paragraph_format.space_before = Pt(12)
            elif "[DIAGRAM" in line or "[රූප සටහන" in line:
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.italic = True
                try: run.font.color.rgb = RGBColor(0, 0, 255)
                except: pass
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph().paragraph_format.space_after = Pt(72)
            else:
                p = doc.add_paragraph(line)
                p.paragraph_format.space_after = Pt(6)
        else:
            if line.startswith("#") or "Subject:" in line:
                p = doc.add_paragraph()
                run = p.add_run(line.replace("#", "").strip())
                run.bold = True
            else:
                doc.add_paragraph(line)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- 6. SIDEBAR ---
with st.sidebar:
    st.title("⚙️ Control Panel")
    if "GEMINI_API_KEY" in st.secrets: api_key = st.secrets["GEMINI_API_KEY"]
    else: api_key = st.text_input("Gemini API Key:", type="password")
    
    st.divider()
    app_mode = st.radio("Select Mode:", ["📝 Exam Paper Generator", "💼 Business Center (Letters/Tutes)"])
    st.divider()
    language = st.radio("Language:", ["Sinhala", "English", "Sinhala + English (Mix)"])
    
    st.divider()
    if app_mode == "📝 Exam Paper Generator":
        st.subheader("📚 Book & Paper Settings")
        deep_search = st.checkbox("🧠 Deep Search (Geometry)", value=False)
        marking_scheme = st.checkbox("✅ Include Marking Scheme", value=True, help="Generate answers at the end of the paper.")
        vision_mode = st.checkbox("🔮 Vision Mode", value=True)
        c1, c2 = st.columns(2)
        with c1: start_p = st.number_input("Start Page", 1, value=1)
        with c2: end_p = st.number_input("End Page", 1, value=50)
    else:
        vision_mode = st.checkbox("🔮 Vision Mode", value=True)
        start_p, end_p = 1, 5

# --- 7. MAIN LOGIC ---
st.title(f"{app_mode}")

if app_mode == "📝 Exam Paper Generator":
    col1, col2 = st.columns(2)
    with col1:
        user_instr = st.text_area("Instructions:", height=100)
        ref_files = st.file_uploader("Reference Paper", accept_multiple_files=True, key="exam_ref")
    with col2:
        src_files = st.file_uploader("Textbooks (PDF)", accept_multiple_files=True, key="exam_src")
        
    if st.button("Generate Paper", type="primary"):
        if not api_key: st.error("Please enter API Key")
        else:
            with st.spinner("Analyzing & Generating..."):
                model = get_working_model(api_key, deep_search)
                content = process_files(src_files, vision_mode, start_p, end_p)
                ref_content = process_files(ref_files, False)
                ref_text = ref_content[0]['data'] if ref_content and ref_content[0]['type'] == 'text' else ""
                
                # Marking Scheme Logic
                answers_instruction = ""
                if marking_scheme:
                    answers_instruction = "Task 3: MARKING SCHEME. After the paper and Typst code, add a section titled 'MARKING SCHEME' and provide brief answers/marking points for all questions."

                lang_instruction = "Mix of Sinhala/English" if language == "Sinhala + English (Mix)" else language
                
                prompt = f"""
                Role: Expert Exam Setter. Language: {lang_instruction}.
                
                Task 1: EXAM PAPER CONTENT
                - Follow Reference Structure Exactly.
                - Rules: Unicode Sinhala, Linear Math, [DIAGRAM] placeholders.
                
                Task 2: TYPST CODE
                - Separator: "### TYPST START ###".
                - Draw diagrams using 'cetz'.
                
                {answers_instruction}
                
                User Instructions: {user_instr}
                Reference: {ref_text[:2500]}
                """
                
                res = call_gemini(api_key, model, prompt, content)
                
                if "### TYPST START ###" in res:
                    parts = res.split("### TYPST START ###")
                    st.session_state.generated_content = parts[0].strip()
                    st.session_state.typst_code = parts[1].strip()
                    # Handle Marking Scheme if it got stuck in Typst part (rare but possible)
                    if "MARKING SCHEME" in parts[1]:
                        sub_parts = parts[1].split("MARKING SCHEME")
                        st.session_state.typst_code = sub_parts[0].strip()
                        st.session_state.generated_content += "\n\nMARKING SCHEME\n" + sub_parts[1].strip()
                else:
                    st.session_state.generated_content = res
                    st.session_state.typst_code = "// No diagram code"
                st.rerun()

    if st.session_state.generated_content:
        st.divider()
        c1, c2 = st.columns(2)
        with c1:
            st.subheader("📄 Word Paper")
            st.download_button("Download .docx", create_docx(st.session_state.generated_content, True), "Paper.docx")
            st.text_area("Preview", st.session_state.generated_content, height=400)
        with c2:
            st.subheader("📐 Typst Diagrams")
            st.download_button("Download .typ", st.session_state.typst_code, "diagrams.typ")
            st.code(st.session_state.typst_code, language="rust")

else:
    # BUSINESS MODE (UNCHANGED BUT ROBUST)
    col1, col2 = st.columns(2)
    with col1:
        doc_type = st.selectbox("Document Type:", ["Official Letter", "Tute / Note", "Report", "Assignment"])
        user_instr = st.text_area("Instructions (Override Image):", height=150)
    with col2:
        src_files = st.file_uploader("Source Docs", accept_multiple_files=True, key="biz_src")
        
    if st.button("Create Document", type="primary"):
        if not api_key: st.error("Please enter API Key")
        else:
            with st.spinner("Drafting..."):
                model = get_working_model(api_key, False)
                content = process_files(src_files, vision_mode, start_p, end_p)
                lang_instruction = "Mix of Sinhala/English" if language == "Sinhala + English (Mix)" else language
                
                prompt = f"""
                Role: Professional Writer. Task: Write a {doc_type}. Language: {lang_instruction}.
                CRITICAL: PRIORITIZE User Instructions below over Image Content.
                User Instructions: {user_instr}
                """
                res = call_gemini(api_key, model, prompt, content)
                st.session_state.generated_content = res
                st.rerun()

    if st.session_state.generated_content:
        st.divider()
        st.subheader(f"📄 Generated {doc_type}")
        st.download_button("Download .docx", create_docx(st.session_state.generated_content, False), "Document.docx")
        st.text_area("Preview", st.session_state.generated_content, height=600)
